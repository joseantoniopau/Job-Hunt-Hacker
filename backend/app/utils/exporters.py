"""Render structured resume dicts to multiple formats.

A structured resume looks like::

    {
      "header": {"name": "...", "email": "...", "phone": "...", "location": "...", "links": [...]},
      "summary": "...",
      "sections": [{"title": "...", "items": [{"text": "..."}]}],
      ...
    }

Markdown and plain-text exports are always available. DOCX requires
``python-docx``; PDF tries ``reportlab`` first, then ``weasyprint``. Each
optional dep is wrapped so a missing package degrades gracefully.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

log = logging.getLogger("jhh.utils.exporters")

try:
    from docx import Document as _DocxDocument  # type: ignore
    from docx.shared import Pt as _DocxPt  # type: ignore
    _DOCX_OK = True
except Exception as _e:  # noqa: BLE001
    _DocxDocument = None  # type: ignore
    _DocxPt = None  # type: ignore
    _DOCX_OK = False
    _DOCX_ERR = str(_e)

try:
    from reportlab.lib.pagesizes import LETTER  # type: ignore
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer  # type: ignore
    from reportlab.lib.units import inch  # type: ignore
    _REPORTLAB_OK = True
except Exception as _e:  # noqa: BLE001
    _REPORTLAB_OK = False
    _REPORTLAB_ERR = str(_e)

try:
    from weasyprint import HTML as _WeasyHTML  # type: ignore
    _WEASY_OK = True
except Exception as _e:  # noqa: BLE001
    _WeasyHTML = None  # type: ignore
    _WEASY_OK = False
    _WEASY_ERR = str(_e)


# ---------- normalization ----------

def _header(resume: dict) -> dict:
    h = resume.get("header") or {}
    if isinstance(h, str):
        return {"name": h}
    if not isinstance(h, dict):
        return {}
    return h


def _sections(resume: dict) -> list[dict]:
    secs = resume.get("sections") or []
    out: list[dict] = []
    for s in secs:
        if not isinstance(s, dict):
            continue
        items = s.get("items") or []
        clean = []
        for it in items:
            if isinstance(it, str):
                clean.append({"text": it})
            elif isinstance(it, dict):
                clean.append(it)
        out.append({"title": s.get("title") or "", "items": clean})
    return out


# ---------- markdown ----------

def to_markdown(resume_dict: dict) -> str:
    h = _header(resume_dict)
    lines: list[str] = []
    name = (h.get("name") or "").strip()
    if name:
        lines.append(f"# {name}")
    contact_bits = [b for b in [
        (h.get("email") or "").strip(),
        (h.get("phone") or "").strip(),
        (h.get("location") or "").strip(),
    ] if b]
    if contact_bits:
        lines.append(" | ".join(contact_bits))
    links = h.get("links") or []
    if isinstance(links, list) and links:
        lines.append(" | ".join(str(l) for l in links if l))

    summary = (resume_dict.get("summary") or "").strip()
    if summary:
        lines.append("")
        lines.append("## Summary")
        lines.append(summary)

    for sec in _sections(resume_dict):
        title = (sec.get("title") or "").strip()
        if title:
            lines.append("")
            lines.append(f"## {title}")
        for it in sec.get("items") or []:
            text = (it.get("text") or "").strip()
            if not text:
                continue
            lines.append(f"- {text}")

    # optional appendices
    gaps = resume_dict.get("gaps") or []
    if isinstance(gaps, list) and gaps:
        lines.append("")
        lines.append("## Gaps (not claimed on resume)")
        for g in gaps:
            lines.append(f"- {g}")

    return "\n".join(lines).strip() + "\n"


# ---------- plain text (ATS safe) ----------

def to_plain_text(resume_dict: dict) -> str:
    h = _header(resume_dict)
    lines: list[str] = []
    name = (h.get("name") or "").strip()
    if name:
        lines.append(name.upper())
    contact_bits = [b for b in [
        (h.get("email") or "").strip(),
        (h.get("phone") or "").strip(),
        (h.get("location") or "").strip(),
    ] if b]
    if contact_bits:
        lines.append(" | ".join(contact_bits))
    links = h.get("links") or []
    if isinstance(links, list) and links:
        lines.append(" | ".join(str(l) for l in links if l))

    summary = (resume_dict.get("summary") or "").strip()
    if summary:
        lines.append("")
        lines.append("SUMMARY")
        lines.append(summary)

    for sec in _sections(resume_dict):
        title = (sec.get("title") or "").strip()
        if title:
            lines.append("")
            lines.append(title.upper())
        for it in sec.get("items") or []:
            text = (it.get("text") or "").strip()
            if not text:
                continue
            # ATS-safe bullet: leading dash, no unicode
            lines.append(f"- {text}")

    return "\n".join(lines).strip() + "\n"


# ---------- docx ----------

def to_docx(resume_dict: dict, out_path: Path) -> Path:
    if not _DOCX_OK:
        raise RuntimeError("python-docx is required for DOCX export")
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = _DocxDocument()  # type: ignore
    # Set default font/size on Normal style
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = _DocxPt(11)  # type: ignore
    except Exception:
        pass

    h = _header(resume_dict)
    name = (h.get("name") or "").strip()
    if name:
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        run.font.size = _DocxPt(18)  # type: ignore

    contact_bits = [b for b in [
        (h.get("email") or "").strip(),
        (h.get("phone") or "").strip(),
        (h.get("location") or "").strip(),
    ] if b]
    if contact_bits:
        doc.add_paragraph(" | ".join(contact_bits))
    links = h.get("links") or []
    if isinstance(links, list) and links:
        doc.add_paragraph(" | ".join(str(l) for l in links if l))

    summary = (resume_dict.get("summary") or "").strip()
    if summary:
        p = doc.add_paragraph()
        run = p.add_run("Summary")
        run.bold = True
        run.font.size = _DocxPt(14)  # type: ignore
        doc.add_paragraph(summary)

    for sec in _sections(resume_dict):
        title = (sec.get("title") or "").strip()
        if title:
            p = doc.add_paragraph()
            run = p.add_run(title)
            run.bold = True
            run.font.size = _DocxPt(14)  # type: ignore
        for it in sec.get("items") or []:
            text = (it.get("text") or "").strip()
            if not text:
                continue
            doc.add_paragraph(text, style=None).paragraph_format.left_indent = _DocxPt(12)  # type: ignore

    doc.save(str(out_path))
    return out_path


# ---------- pdf (best effort) ----------

def _pdf_via_reportlab(resume_dict: dict, out_path: Path) -> Path:
    styles = getSampleStyleSheet()  # type: ignore
    body = styles["BodyText"]
    h1 = ParagraphStyle(  # type: ignore
        "JhhH1", parent=styles["Heading1"], fontSize=18, leading=22, spaceAfter=6
    )
    h2 = ParagraphStyle(  # type: ignore
        "JhhH2", parent=styles["Heading2"], fontSize=13, leading=16, spaceBefore=10, spaceAfter=4
    )
    doc = SimpleDocTemplate(  # type: ignore
        str(out_path), pagesize=LETTER,  # type: ignore
        leftMargin=0.6 * inch, rightMargin=0.6 * inch,  # type: ignore
        topMargin=0.5 * inch, bottomMargin=0.5 * inch,  # type: ignore
    )
    story: list = []
    h = _header(resume_dict)
    name = (h.get("name") or "").strip()
    if name:
        story.append(Paragraph(name, h1))  # type: ignore
    bits = [b for b in [h.get("email"), h.get("phone"), h.get("location")] if b]
    if bits:
        story.append(Paragraph(" | ".join(bits), body))  # type: ignore
    links = h.get("links") or []
    if links:
        story.append(Paragraph(" | ".join(str(l) for l in links if l), body))  # type: ignore

    summary = (resume_dict.get("summary") or "").strip()
    if summary:
        story.append(Paragraph("Summary", h2))  # type: ignore
        story.append(Paragraph(summary, body))  # type: ignore

    for sec in _sections(resume_dict):
        title = (sec.get("title") or "").strip()
        if title:
            story.append(Paragraph(title, h2))  # type: ignore
        for it in sec.get("items") or []:
            text = (it.get("text") or "").strip()
            if not text:
                continue
            story.append(Paragraph("&bull; " + _xml_escape(text), body))  # type: ignore
        story.append(Spacer(1, 4))  # type: ignore

    doc.build(story)
    return out_path


def _xml_escape(s: str) -> str:
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))


def _pdf_via_weasy(resume_dict: dict, out_path: Path) -> Path:
    # Build a small HTML doc and let weasyprint render it.
    h = _header(resume_dict)
    parts: list[str] = ["<html><head><meta charset='utf-8'><style>",
                        "body{font-family:Helvetica,Arial,sans-serif;font-size:11pt;line-height:1.4;margin:0.6in;}",
                        "h1{font-size:18pt;margin:0 0 4pt 0;} h2{font-size:13pt;margin:12pt 0 4pt 0;}",
                        "ul{margin:0 0 0 16pt;padding:0;} li{margin:0 0 2pt 0;}",
                        ".meta{color:#444;}",
                        "</style></head><body>"]
    name = (h.get("name") or "").strip()
    if name:
        parts.append(f"<h1>{_xml_escape(name)}</h1>")
    bits = [b for b in [h.get("email"), h.get("phone"), h.get("location")] if b]
    if bits:
        parts.append("<div class='meta'>" + _xml_escape(" | ".join(bits)) + "</div>")
    links = h.get("links") or []
    if links:
        parts.append("<div class='meta'>" + _xml_escape(" | ".join(str(l) for l in links if l)) + "</div>")
    summary = (resume_dict.get("summary") or "").strip()
    if summary:
        parts.append("<h2>Summary</h2>")
        parts.append("<div>" + _xml_escape(summary) + "</div>")
    for sec in _sections(resume_dict):
        title = (sec.get("title") or "").strip()
        if title:
            parts.append(f"<h2>{_xml_escape(title)}</h2>")
        items = sec.get("items") or []
        if items:
            parts.append("<ul>")
            for it in items:
                t = (it.get("text") or "").strip()
                if not t:
                    continue
                parts.append("<li>" + _xml_escape(t) + "</li>")
            parts.append("</ul>")
    parts.append("</body></html>")
    html = "".join(parts)
    _WeasyHTML(string=html).write_pdf(str(out_path))  # type: ignore
    return out_path


def to_pdf(resume_dict: dict, out_path: Path) -> Path | None:
    """Best-effort PDF export. Returns the output path or None if no
    PDF backend is available.
    """
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    if _REPORTLAB_OK:
        try:
            return _pdf_via_reportlab(resume_dict, out_path)
        except Exception as e:  # noqa: BLE001
            log.warning("reportlab pdf export failed: %s", e)
    if _WEASY_OK:
        try:
            return _pdf_via_weasy(resume_dict, out_path)
        except Exception as e:  # noqa: BLE001
            log.warning("weasyprint pdf export failed: %s", e)
    log.info("PDF export requires reportlab or weasyprint; skipping")
    return None


__all__ = ["to_markdown", "to_plain_text", "to_docx", "to_pdf"]
