"""Parse uploaded documents into plain text + metadata.

Defensive imports: each optional dependency is wrapped so a missing package
degrades to a clear runtime error only when the relevant code path is taken.
"""
from __future__ import annotations

import logging
import mimetypes
from pathlib import Path
from typing import Any

log = logging.getLogger("jhh.evidence")

try:
    from pdfminer.high_level import extract_text as _pdf_extract_text  # type: ignore
    _PDF_OK = True
except Exception as _e:  # noqa: BLE001
    _pdf_extract_text = None  # type: ignore
    _PDF_OK = False
    _PDF_ERR = str(_e)

try:
    import docx as _docx  # type: ignore
    _DOCX_OK = True
except Exception as _e:  # noqa: BLE001
    _docx = None  # type: ignore
    _DOCX_OK = False
    _DOCX_ERR = str(_e)

try:
    from bs4 import BeautifulSoup as _BS  # type: ignore
    _BS_OK = True
except Exception as _e:  # noqa: BLE001
    _BS = None  # type: ignore
    _BS_OK = False
    _BS_ERR = str(_e)


def _read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:  # noqa: BLE001
        log.warning("text read failed for %s: %s", path, e)
        return ""


def _parse_pdf(path: Path) -> str:
    if not _PDF_OK:
        raise RuntimeError("install pdfminer.six to parse PDF files")
    try:
        return _pdf_extract_text(str(path)) or ""  # type: ignore
    except Exception as e:  # noqa: BLE001
        log.warning("pdf parse failed for %s: %s", path, e)
        return ""


def _parse_docx(path: Path) -> str:
    if not _DOCX_OK:
        raise RuntimeError("install python-docx to parse DOCX files")
    try:
        doc = _docx.Document(str(path))  # type: ignore
        parts = [p.text for p in doc.paragraphs]
        # Also pull text from any tables.
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(p for p in parts if p)
    except Exception as e:  # noqa: BLE001
        log.warning("docx parse failed for %s: %s", path, e)
        return ""


def _parse_html(path: Path) -> str:
    raw = _read_text(path)
    if not _BS_OK:
        # crude strip
        import re
        return re.sub(r"<[^>]+>", " ", raw)
    try:
        soup = _BS(raw, "html.parser")  # type: ignore
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        return soup.get_text("\n").strip()
    except Exception as e:  # noqa: BLE001
        log.warning("html parse failed for %s: %s", path, e)
        return raw


def parse_file(path: Path) -> dict[str, Any]:
    """Return ``{"text": str, "metadata": dict}``.

    Supports PDF, DOCX, TXT, MD, HTML. Unknown types fall back to a UTF-8
    read so callers always get something usable.
    """
    path = Path(path)
    suffix = path.suffix.lower().lstrip(".")
    metadata: dict[str, Any] = {
        "filename": path.name,
        "suffix": suffix,
        "size_bytes": path.stat().st_size if path.exists() else 0,
        "mime": mimetypes.guess_type(str(path))[0] or "",
    }

    text = ""
    try:
        if suffix == "pdf":
            text = _parse_pdf(path)
            metadata["parser"] = "pdfminer"
        elif suffix in ("docx",):
            text = _parse_docx(path)
            metadata["parser"] = "python-docx"
        elif suffix in ("html", "htm"):
            text = _parse_html(path)
            metadata["parser"] = "beautifulsoup"
        elif suffix in ("txt", "md", "markdown", "rst", "log", ""):
            text = _read_text(path)
            metadata["parser"] = "text"
        else:
            # Unknown — try text. Caller can decide what to do with garbage.
            text = _read_text(path)
            metadata["parser"] = "text-fallback"
    except RuntimeError:
        raise
    except Exception as e:  # noqa: BLE001
        log.warning("parse_file unexpected failure for %s: %s", path, e)
        text = _read_text(path)
        metadata["parser"] = "error-fallback"
        metadata["error"] = str(e)

    text = (text or "").strip()
    metadata["char_count"] = len(text)
    return {"text": text, "metadata": metadata}
